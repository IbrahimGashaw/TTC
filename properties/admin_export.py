"""Generic admin export utilities for CSV, Excel, Word, and PDF."""
from __future__ import annotations

import csv
from dataclasses import dataclass
from datetime import datetime
from io import BytesIO
from typing import Any, Callable

from django.db.models import QuerySet
from django.http import HttpResponse
from django.utils import timezone
from django.utils.text import slugify
from django.utils.translation import gettext as _


@dataclass(frozen=True)
class ExportColumn:
    key: str
    label: str
    getter: Callable[[Any], Any] | None = None


@dataclass(frozen=True)
class ExportConfig:
    title: str
    filename_prefix: str
    columns: tuple[ExportColumn, ...]


def _cell_value(obj, column: ExportColumn):
    if column.getter:
        value = column.getter(obj)
    else:
        value = getattr(obj, column.key, '')
        if callable(value) and not isinstance(value, (str, bytes, int, float, bool)):
            value = value()
    if value is None:
        return ''
    if hasattr(value, 'strftime'):
        return timezone.localtime(value).strftime('%Y-%m-%d %H:%M')
    return value


def _header_labels(config: ExportConfig):
    return [str(column.label) for column in config.columns]


def _data_rows(queryset: QuerySet, config: ExportConfig):
    rows = []
    for obj in queryset:
        rows.append([_cell_value(obj, column) for column in config.columns])
    return rows


def _build_filename(config: ExportConfig, extension: str) -> str:
    slug = slugify(config.filename_prefix) or 'export'
    date_stamp = datetime.now().strftime('%Y%m%d')
    return f'{slug}-{date_stamp}.{extension}'


def _summary_line(config: ExportConfig, count: int) -> str:
    exported_at = timezone.localtime(timezone.now()).strftime('%Y-%m-%d %H:%M')
    return f'{config.title} | {_("Total records")}: {count} | {_("Exported")}: {exported_at}'


def export_csv(queryset: QuerySet, config: ExportConfig):
    response = HttpResponse(content_type='text/csv; charset=utf-8')
    response['Content-Disposition'] = f'attachment; filename="{_build_filename(config, "csv")}"'
    response.write('\ufeff')
    writer = csv.writer(response)
    writer.writerow([str(_('Export')), _summary_line(config, queryset.count())])
    writer.writerow([])
    writer.writerow(_header_labels(config))
    writer.writerows(_data_rows(queryset, config))
    return response


def export_xlsx(queryset: QuerySet, config: ExportConfig):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = slugify(config.title)[:31] or 'Export'

    header_fill = PatternFill(start_color='1E5A96', end_color='1E5A96', fill_type='solid')
    header_font = Font(color='FFFFFF', bold=True)

    sheet.append([str(_('Export')), _summary_line(config, queryset.count())])
    sheet.append([])

    labels = _header_labels(config)
    sheet.append(labels)
    for cell in sheet[sheet.max_row]:
        cell.fill = header_fill
        cell.font = header_font

    for row in _data_rows(queryset, config):
        sheet.append(row)

    for column in sheet.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            if cell.value is not None:
                max_length = max(max_length, len(str(cell.value)))
        sheet.column_dimensions[column_letter].width = min(max_length + 2, 60)

    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = f'attachment; filename="{_build_filename(config, "xlsx")}"'
    return response


def export_docx(queryset: QuerySet, config: ExportConfig):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    document = Document()
    title = document.add_heading(str(config.title), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(_summary_line(config, queryset.count()))
    document.add_paragraph('')

    labels = _header_labels(config)
    table = document.add_table(rows=1, cols=len(labels))
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for index, label in enumerate(labels):
        header_cells[index].text = label
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for row in _data_rows(queryset, config):
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            row_cells[index].text = str(value)

    for section in document.sections:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{_build_filename(config, "docx")}"'
    return response


def export_pdf(queryset: QuerySet, config: ExportConfig):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=0.4 * inch,
        rightMargin=0.4 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
    )
    styles = getSampleStyleSheet()
    story = [
        Paragraph(str(config.title), styles['Title']),
        Paragraph(_summary_line(config, queryset.count()), styles['Normal']),
        Spacer(1, 12),
    ]

    table_data = [_header_labels(config)] + _data_rows(queryset, config)
    table = Table(table_data, repeatRows=1)
    table.setStyle(
        TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E5A96')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 7),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F4F8FC')]),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ])
    )
    story.append(table)
    document.build(story)
    buffer.seek(0)

    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{_build_filename(config, "pdf")}"'
    return response


EXPORTERS = {
    'csv': export_csv,
    'xlsx': export_xlsx,
    'docx': export_docx,
    'pdf': export_pdf,
}


def export_queryset(queryset: QuerySet, config: ExportConfig, file_format: str):
    exporter = EXPORTERS.get(file_format)
    if exporter is None:
        raise ValueError(f'Unsupported export format: {file_format}')
    return exporter(queryset, config)


def _yes_no(value):
    return _('Yes') if value else _('No')


def _local_datetime(value):
    if not value:
        return ''
    return timezone.localtime(value).strftime('%Y-%m-%d %H:%M')


CONTACT_EXPORT = ExportConfig(
    title=_('Contact Messages'),
    filename_prefix='contact-messages',
    columns=(
        ExportColumn('id', _('ID')),
        ExportColumn('name', _('First Name')),
        ExportColumn('last_name', _('Last Name')),
        ExportColumn('email', _('Email')),
        ExportColumn('phone', _('Phone')),
        ExportColumn('subject', _('Subject'), lambda o: o.get_subject_display()),
        ExportColumn('message', _('Message')),
        ExportColumn('gdpr_consent', _('Consent'), lambda o: _yes_no(o.gdpr_consent)),
        ExportColumn('created_at', _('Submitted At'), lambda o: _local_datetime(o.created_at)),
    ),
)

EVENT_REGISTRATION_EXPORT = ExportConfig(
    title=_('Event Registrations'),
    filename_prefix='event-registrations',
    columns=(
        ExportColumn('event', _('Event'), lambda o: o.event.title if o.event_id else ''),
        ExportColumn('full_name', _('Full Name')),
        ExportColumn('email', _('Email')),
        ExportColumn('phone', _('Phone')),
        ExportColumn('organization', _('Organization')),
        ExportColumn('job_title', _('Job Title')),
        ExportColumn('status', _('Status'), lambda o: o.get_status_display()),
        ExportColumn('notes', _('Notes')),
        ExportColumn('gdpr_consent', _('Consent'), lambda o: _yes_no(o.gdpr_consent)),
        ExportColumn('created_at', _('Registered At'), lambda o: _local_datetime(o.created_at)),
    ),
)

SERVICE_EXPORT = ExportConfig(
    title=_('Services'),
    filename_prefix='services',
    columns=(
        ExportColumn('title', _('Title')),
        ExportColumn('category', _('Category'), lambda o: o.get_category_display()),
        ExportColumn('short_description', _('Short Description')),
        ExportColumn('is_featured', _('Featured'), lambda o: _yes_no(o.is_featured)),
        ExportColumn('is_active', _('Active'), lambda o: _yes_no(o.is_active)),
        ExportColumn('order', _('Order')),
    ),
)

TRAINING_EVENT_EXPORT = ExportConfig(
    title=_('Training Events'),
    filename_prefix='training-events',
    columns=(
        ExportColumn('title', _('Title')),
        ExportColumn('event_type', _('Type'), lambda o: o.get_event_type_display()),
        ExportColumn('location', _('Location')),
        ExportColumn('start_date', _('Start Date'), lambda o: _local_datetime(o.start_date)),
        ExportColumn('end_date', _('End Date'), lambda o: _local_datetime(o.end_date)),
        ExportColumn('is_published', _('Published'), lambda o: _yes_no(o.is_published)),
        ExportColumn('max_participants', _('Max Participants')),
    ),
)

TEAM_MEMBER_EXPORT = ExportConfig(
    title=_('Team Members'),
    filename_prefix='team-members',
    columns=(
        ExportColumn('name', _('Name')),
        ExportColumn('title', _('Title')),
        ExportColumn('role', _('Role'), lambda o: o.get_role_display()),
        ExportColumn('email', _('Email')),
        ExportColumn('phone', _('Phone')),
        ExportColumn('years_experience', _('Years Experience')),
        ExportColumn('is_founder', _('Founder'), lambda o: _yes_no(o.is_founder)),
    ),
)

TESTIMONIAL_EXPORT = ExportConfig(
    title=_('Testimonials'),
    filename_prefix='testimonials',
    columns=(
        ExportColumn('client_name', _('Client')),
        ExportColumn('organization', _('Organization')),
        ExportColumn('content', _('Content')),
        ExportColumn('rating', _('Rating')),
        ExportColumn('is_featured', _('Featured'), lambda o: _yes_no(o.is_featured)),
        ExportColumn('is_active', _('Active'), lambda o: _yes_no(o.is_active)),
    ),
)

CASE_STUDY_EXPORT = ExportConfig(
    title=_('Case Studies'),
    filename_prefix='case-studies',
    columns=(
        ExportColumn('title', _('Title')),
        ExportColumn('client_name', _('Client')),
        ExportColumn('industry', _('Industry')),
        ExportColumn('excerpt', _('Excerpt')),
        ExportColumn('is_featured', _('Featured'), lambda o: _yes_no(o.is_featured)),
        ExportColumn('is_published', _('Published'), lambda o: _yes_no(o.is_published)),
    ),
)

PARTNER_EXPORT = ExportConfig(
    title=_('Partners'),
    filename_prefix='partners',
    columns=(
        ExportColumn('name', _('Name')),
        ExportColumn('website', _('Website')),
        ExportColumn('description', _('Description')),
        ExportColumn('is_active', _('Active'), lambda o: _yes_no(o.is_active)),
        ExportColumn('order', _('Order')),
    ),
)

VACANCY_EXPORT = ExportConfig(
    title=_('Vacancies'),
    filename_prefix='vacancies',
    columns=(
        ExportColumn('title', _('Title')),
        ExportColumn('department', _('Department')),
        ExportColumn('location', _('Location')),
        ExportColumn('employment_type', _('Type'), lambda o: o.get_employment_type_display()),
        ExportColumn('application_open_at', _('Opens'), lambda o: _local_datetime(o.application_open_at)),
        ExportColumn('application_deadline', _('Deadline'), lambda o: _local_datetime(o.application_deadline)),
        ExportColumn('max_applications', _('Max Applications')),
        ExportColumn('is_published', _('Published'), lambda o: _yes_no(o.is_published)),
        ExportColumn('is_featured', _('Featured'), lambda o: _yes_no(o.is_featured)),
    ),
)

VACANCY_APPLICATION_EXPORT = ExportConfig(
    title=_('Vacancy Applications'),
    filename_prefix='vacancy-applications',
    columns=(
        ExportColumn('vacancy', _('Vacancy'), lambda o: o.vacancy.title if o.vacancy_id else ''),
        ExportColumn('full_name', _('Full Name')),
        ExportColumn('email', _('Email')),
        ExportColumn('phone', _('Phone')),
        ExportColumn('education', _('Education')),
        ExportColumn('years_experience', _('Years of Experience')),
        ExportColumn('status', _('Status'), lambda o: o.get_status_display()),
        ExportColumn('cover_letter', _('Cover Letter')),
        ExportColumn('resume', _('Resume'), lambda o: o.resume.name.split('/')[-1] if o.resume else ''),
        ExportColumn('gdpr_consent', _('Consent'), lambda o: _yes_no(o.gdpr_consent)),
        ExportColumn('created_at', _('Applied At'), lambda o: _local_datetime(o.created_at)),
    ),
)

MILESTONE_EXPORT = ExportConfig(
    title=_('Milestones'),
    filename_prefix='milestones',
    columns=(
        ExportColumn('year', _('Year')),
        ExportColumn('title', _('Title')),
        ExportColumn('description', _('Description')),
        ExportColumn('order', _('Order')),
    ),
)

WORKING_SCHEDULE_EXPORT = ExportConfig(
    title=_('Working Schedule'),
    filename_prefix='working-schedule',
    columns=(
        ExportColumn('day_of_week', _('Day'), lambda o: o.get_day_of_week_display()),
        ExportColumn('open_time', _('Opens')),
        ExportColumn('close_time', _('Closes')),
        ExportColumn('is_closed', _('Closed'), lambda o: _yes_no(o.is_closed)),
    ),
)

BLOG_POST_EXPORT = ExportConfig(
    title=_('Blog Posts'),
    filename_prefix='blog-posts',
    columns=(
        ExportColumn('title', _('Title')),
        ExportColumn('category', _('Category'), lambda o: o.category.name if o.category_id else ''),
        ExportColumn('author', _('Author'), lambda o: o.author.get_full_name() or o.author.username),
        ExportColumn('published', _('Published'), lambda o: _yes_no(o.published)),
        ExportColumn('created_at', _('Created At'), lambda o: _local_datetime(o.created_at)),
    ),
)


def make_export_action(config: ExportConfig, file_format: str, label: str):
    def action(modeladmin, request, queryset):
        if not queryset.exists():
            queryset = modeladmin.model.objects.all()
        queryset = queryset.order_by('pk')
        return export_queryset(queryset, config, file_format)

    action.short_description = label
    action.__name__ = f'export_{config.filename_prefix}_{file_format}'
    return action


def export_admin_actions(config: ExportConfig):
    return [
        make_export_action(config, 'xlsx', _('Export to Excel (.xlsx)')),
        make_export_action(config, 'csv', _('Export to CSV (.csv)')),
        make_export_action(config, 'docx', _('Export to Word (.docx)')),
        make_export_action(config, 'pdf', _('Export to PDF (.pdf)')),
    ]


class ExportMixin:
    """Add export actions to a ModelAdmin subclass."""

    export_config: ExportConfig | None = None

    def get_actions(self, request):
        actions = super().get_actions(request)
        if self.export_config:
            for action in export_admin_actions(self.export_config):
                actions[action.__name__] = (action, action.__name__, action.short_description)
        return actions


ALL_EXPORT_CONFIGS: tuple[tuple[str, str, ExportConfig], ...] = (
    ('properties', 'Contact', CONTACT_EXPORT),
    ('properties', 'EventRegistration', EVENT_REGISTRATION_EXPORT),
    ('properties', 'Vacancy', VACANCY_EXPORT),
    ('properties', 'VacancyApplication', VACANCY_APPLICATION_EXPORT),
    ('properties', 'Service', SERVICE_EXPORT),
    ('properties', 'TrainingEvent', TRAINING_EVENT_EXPORT),
    ('properties', 'TeamMember', TEAM_MEMBER_EXPORT),
    ('properties', 'Testimonial', TESTIMONIAL_EXPORT),
    ('properties', 'CaseStudy', CASE_STUDY_EXPORT),
    ('properties', 'Partner', PARTNER_EXPORT),
    ('properties', 'Milestone', MILESTONE_EXPORT),
    ('properties', 'WorkingSchedule', WORKING_SCHEDULE_EXPORT),
    ('blog', 'BlogPost', BLOG_POST_EXPORT),
)


def build_all_data_zip():
    """Build a ZIP archive containing Excel exports for all configured models."""
    import zipfile
    from django.apps import apps

    buffer = BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as archive:
        for app_label, model_name, config in ALL_EXPORT_CONFIGS:
            model = apps.get_model(app_label, model_name)
            queryset = model.objects.all().order_by('pk')
            if not queryset.exists():
                continue
            response = export_xlsx(queryset, config)
            filename = response['Content-Disposition'].split('filename=')[1].strip('"')
            archive.writestr(filename, response.content)
    buffer.seek(0)
    return buffer


def export_all_data_zip_response():
    from django.http import HttpResponse

    date_stamp = datetime.now().strftime('%Y%m%d')
    response = HttpResponse(
        build_all_data_zip().getvalue(),
        content_type='application/zip',
    )
    response['Content-Disposition'] = f'attachment; filename="ttcs-export-{date_stamp}.zip"'
    return response
