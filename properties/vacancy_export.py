"""Export vacancy applications to CSV, Excel, Word, and PDF."""
from __future__ import annotations

import csv
from datetime import datetime
from io import BytesIO

from django.http import HttpResponse
from django.utils import timezone
from django.utils.text import slugify
from django.utils.translation import gettext as _

EXPORT_COLUMNS = [
    ('full_name', _('Full Name')),
    ('email', _('Email')),
    ('phone', _('Phone')),
    ('education', _('Education')),
    ('years_experience', _('Years of Experience')),
    ('status', _('Status')),
    ('cover_letter', _('Cover Letter')),
    ('resume', _('Resume')),
    ('applied_at', _('Applied At')),
]


def _status_label(application):
    return application.get_status_display()


def _applied_at(application):
    dt = timezone.localtime(application.created_at)
    return dt.strftime('%Y-%m-%d %H:%M')


def _resume_label(application):
    if application.resume:
        return application.resume.name.split('/')[-1]
    return ''


def _application_row(application):
    return {
        'full_name': application.full_name,
        'email': application.email,
        'phone': application.phone,
        'education': application.education,
        'years_experience': application.years_experience if application.years_experience is not None else '',
        'status': _status_label(application),
        'cover_letter': application.cover_letter,
        'resume': _resume_label(application),
        'applied_at': _applied_at(application),
    }


def _header_labels():
    return [str(label) for _, label in EXPORT_COLUMNS]


def _data_rows(queryset):
    rows = []
    for application in queryset.select_related('vacancy').order_by('created_at'):
        row = _application_row(application)
        rows.append([row[key] for key, _ in EXPORT_COLUMNS])
    return rows


def _build_filename(vacancy, extension):
    if vacancy:
        slug = slugify(vacancy.title) or f'vacancy-{vacancy.pk}'
    else:
        slug = 'all-vacancies'
    date_stamp = datetime.now().strftime('%Y%m%d')
    return f'applicants-{slug}-{date_stamp}.{extension}'


def _vacancy_summary(vacancy):
    if not vacancy:
        return _('All selected applications')
    parts = [vacancy.title, vacancy.get_employment_type_display()]
    if vacancy.department:
        parts.append(vacancy.department)
    if vacancy.location:
        parts.append(vacancy.location)
    return ' | '.join(parts)


def export_csv(queryset, vacancy=None):
    response = HttpResponse(content_type='text/csv; charset=utf-8')
    response['Content-Disposition'] = f'attachment; filename="{_build_filename(vacancy, "csv")}"'
    response.write('\ufeff')
    writer = csv.writer(response)
    writer.writerow([_('Vacancy'), _vacancy_summary(vacancy)])
    writer.writerow([_('Exported'), timezone.localtime(timezone.now()).strftime('%Y-%m-%d %H:%M')])
    writer.writerow([])
    writer.writerow(_header_labels())
    writer.writerows(_data_rows(queryset))
    return response


def export_xlsx(queryset, vacancy=None):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = 'Applicants'

    header_fill = PatternFill(start_color='1E5A96', end_color='1E5A96', fill_type='solid')
    header_font = Font(color='FFFFFF', bold=True)

    sheet.append([str(_('Vacancy')), _vacancy_summary(vacancy)])
    sheet.append([str(_('Exported')), timezone.localtime(timezone.now()).strftime('%Y-%m-%d %H:%M')])
    sheet.append([str(_('Total applicants')), queryset.count()])
    sheet.append([])

    labels = _header_labels()
    sheet.append(labels)
    for cell in sheet[sheet.max_row]:
        cell.fill = header_fill
        cell.font = header_font

    for row in _data_rows(queryset):
        sheet.append(row)

    for column in sheet.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        sheet.column_dimensions[column_letter].width = min(max_length + 2, 50)

    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = f'attachment; filename="{_build_filename(vacancy, "xlsx")}"'
    return response


def export_docx(queryset, vacancy=None):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    document = Document()
    title = document.add_heading(str(_('Vacancy Applicants')), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(f'{_("Vacancy")}: {_vacancy_summary(vacancy)}')
    document.add_paragraph(
        f'{_("Exported")}: {timezone.localtime(timezone.now()).strftime("%Y-%m-%d %H:%M")}'
    )
    document.add_paragraph(f'{_("Total applicants")}: {queryset.count()}')
    document.add_paragraph('')

    labels = _header_labels()
    table = document.add_table(rows=1, cols=len(labels))
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for index, label in enumerate(labels):
        header_cells[index].text = label
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for row in _data_rows(queryset):
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            row_cells[index].text = str(value)

    for section in document.sections:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{_build_filename(vacancy, "docx")}"'
    return response


def export_pdf(queryset, vacancy=None):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=0.5 * inch,
        rightMargin=0.5 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
    )
    styles = getSampleStyleSheet()
    story = [
        Paragraph(str(_('Vacancy Applicants')), styles['Title']),
        Paragraph(f'{_("Vacancy")}: {_vacancy_summary(vacancy)}', styles['Normal']),
        Paragraph(
            f'{_("Exported")}: {timezone.localtime(timezone.now()).strftime("%Y-%m-%d %H:%M")} | '
            f'{_("Total applicants")}: {queryset.count()}',
            styles['Normal'],
        ),
        Spacer(1, 12),
    ]

    table_data = [_header_labels()] + _data_rows(queryset)
    table = Table(table_data, repeatRows=1)
    table.setStyle(
        TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E5A96')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F4F8FC')]),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ])
    )
    story.append(table)
    document.build(story)
    buffer.seek(0)

    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{_build_filename(vacancy, "pdf")}"'
    return response


EXPORTERS = {
    'csv': export_csv,
    'xlsx': export_xlsx,
    'docx': export_docx,
    'pdf': export_pdf,
}


def export_applications(queryset, file_format, vacancy=None):
    exporter = EXPORTERS.get(file_format)
    if exporter is None:
        raise ValueError(f'Unsupported export format: {file_format}')
    return exporter(queryset, vacancy=vacancy)
