"""Export event registration participants to CSV, Excel, Word, and PDF."""
from __future__ import annotations

import csv
from datetime import datetime
from io import BytesIO

from django.http import HttpResponse
from django.utils import timezone
from django.utils.text import slugify
from django.utils.translation import gettext as _

EXPORT_COLUMNS = [
    ('full_name', _('Full Name')),
    ('email', _('Email')),
    ('phone', _('Phone')),
    ('organization', _('Organization')),
    ('job_title', _('Job Title')),
    ('status', _('Status')),
    ('notes', _('Notes')),
    ('registered_at', _('Registered At')),
]


def _status_label(registration):
    return registration.get_status_display()


def _registered_at(registration):
    dt = timezone.localtime(registration.created_at)
    return dt.strftime('%Y-%m-%d %H:%M')


def _registration_row(registration):
    return {
        'full_name': registration.full_name,
        'email': registration.email,
        'phone': registration.phone,
        'organization': registration.organization,
        'job_title': registration.job_title,
        'status': _status_label(registration),
        'notes': registration.notes,
        'registered_at': _registered_at(registration),
    }


def _header_labels():
    return [str(label) for _, label in EXPORT_COLUMNS]


def _data_rows(queryset):
    rows = []
    for registration in queryset.select_related('event').order_by('created_at'):
        row = _registration_row(registration)
        rows.append([row[key] for key, _ in EXPORT_COLUMNS])
    return rows


def _build_filename(event, extension):
    if event:
        slug = slugify(event.title) or f'event-{event.pk}'
    else:
        slug = 'all-events'
    date_stamp = datetime.now().strftime('%Y%m%d')
    return f'participants-{slug}-{date_stamp}.{extension}'


def _event_summary(event):
    if not event:
        return _('All selected registrations')
    return (
        f'{event.title} | {event.get_event_type_display()} | '
        f'{timezone.localtime(event.start_date).strftime("%Y-%m-%d %H:%M")} | '
        f'{event.location}'
    )


def export_csv(queryset, event=None):
    response = HttpResponse(content_type='text/csv; charset=utf-8')
    response['Content-Disposition'] = f'attachment; filename="{_build_filename(event, "csv")}"'
    response.write('\ufeff')  # UTF-8 BOM for Excel
    writer = csv.writer(response)
    writer.writerow([_('Event'), _event_summary(event)])
    writer.writerow([_('Exported'), timezone.localtime(timezone.now()).strftime('%Y-%m-%d %H:%M')])
    writer.writerow([])
    writer.writerow(_header_labels())
    writer.writerows(_data_rows(queryset))
    return response


def export_xlsx(queryset, event=None):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = 'Participants'

    header_fill = PatternFill(start_color='1E5A96', end_color='1E5A96', fill_type='solid')
    header_font = Font(color='FFFFFF', bold=True)

    sheet.append([str(_('Event')), _event_summary(event)])
    sheet.append([str(_('Exported')), timezone.localtime(timezone.now()).strftime('%Y-%m-%d %H:%M')])
    sheet.append([str(_('Total participants')), queryset.count()])
    sheet.append([])

    labels = _header_labels()
    sheet.append(labels)
    for cell in sheet[sheet.max_row]:
        cell.fill = header_fill
        cell.font = header_font

    for row in _data_rows(queryset):
        sheet.append(row)

    for column in sheet.columns:
        max_length = 0
        column_letter = column[0].column_letter
        for cell in column:
            if cell.value:
                max_length = max(max_length, len(str(cell.value)))
        sheet.column_dimensions[column_letter].width = min(max_length + 2, 50)

    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = f'attachment; filename="{_build_filename(event, "xlsx")}"'
    return response


def export_docx(queryset, event=None):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    document = Document()
    title = document.add_heading(str(_('Event Participants')), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(f'{_("Event")}: {_event_summary(event)}')
    document.add_paragraph(
        f'{_("Exported")}: {timezone.localtime(timezone.now()).strftime("%Y-%m-%d %H:%M")}'
    )
    document.add_paragraph(f'{_("Total participants")}: {queryset.count()}')
    document.add_paragraph('')

    labels = _header_labels()
    table = document.add_table(rows=1, cols=len(labels))
    table.style = 'Table Grid'
    header_cells = table.rows[0].cells
    for index, label in enumerate(labels):
        header_cells[index].text = label
        for paragraph in header_cells[index].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    for row in _data_rows(queryset):
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            row_cells[index].text = str(value)

    for section in document.sections:
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{_build_filename(event, "docx")}"'
    return response


def export_pdf(queryset, event=None):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=0.5 * inch,
        rightMargin=0.5 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
    )
    styles = getSampleStyleSheet()
    story = [
        Paragraph(str(_('Event Participants')), styles['Title']),
        Paragraph(f'{_("Event")}: {_event_summary(event)}', styles['Normal']),
        Paragraph(
            f'{_("Exported")}: {timezone.localtime(timezone.now()).strftime("%Y-%m-%d %H:%M")} | '
            f'{_("Total participants")}: {queryset.count()}',
            styles['Normal'],
        ),
        Spacer(1, 12),
    ]

    table_data = [_header_labels()] + _data_rows(queryset)
    table = Table(table_data, repeatRows=1)
    table.setStyle(
        TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E5A96')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F4F8FC')]),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
        ])
    )
    story.append(table)
    document.build(story)
    buffer.seek(0)

    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{_build_filename(event, "pdf")}"'
    return response


EXPORTERS = {
    'csv': export_csv,
    'xlsx': export_xlsx,
    'docx': export_docx,
    'pdf': export_pdf,
}


def export_registrations(queryset, file_format, event=None):
    exporter = EXPORTERS.get(file_format)
    if exporter is None:
        raise ValueError(f'Unsupported export format: {file_format}')
    return exporter(queryset, event=event)
